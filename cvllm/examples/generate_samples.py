from __future__ import annotations
from docx import Document
import os

SAMPLES = [
    {
        "name": "sample_resume1.docx",
        "full_name": "Alex Johnson",
        "email": "alex.johnson@example.com",
        "phone": "+1-555-0101",
        "location": "San Francisco, CA",
        "summary": "Software engineer with 5+ years experience building backend services and APIs.",
        "skills": ["Python", "FastAPI", "PostgreSQL", "Docker", "AWS"],
        "experience": [
            (
                "Backend Engineer, Acme Corp (2021–Present)",
                [
                    "Designed and shipped FastAPI services handling 20k RPS with 99.9% availability.",
                    "Optimized PostgreSQL queries reducing p95 latency by 30%.",
                    "Implemented CI/CD on GitHub Actions and containerized services with Docker.",
                ],
            ),
            (
                "Software Engineer, BetaSoft (2019–2021)",
                [
                    "Built internal ETL pipelines in Python feeding analytics dashboards.",
                    "Integrated OAuth and SSO into customer portal; reduced login issues by 40%.",
                ],
            ),
        ],
        "education": [
            "B.S. in Computer Science, State University (2015–2019)",
        ],
    },
    {
        "name": "sample_resume2.docx",
        "full_name": "Priya Singh",
        "email": "priya.singh@example.com",
        "phone": "+91-98765-43210",
        "location": "Bengaluru, IN",
        "summary": "Data scientist experienced in NLP and retrieval systems; shipped ML features to production.",
        "skills": ["Python", "Pandas", "scikit-learn", "PyTorch", "LlamaIndex", "LangChain"],
        "experience": [
            (
                "Data Scientist, InsightAI (2022–Present)",
                [
                    "Developed semantic search with sentence-transformers improving CTR by 22%.",
                    "Deployed LlamaIndex-based retrieval service and integrated rerankers.",
                    "Built evaluation harness and A/B tests for ranking models.",
                ],
            ),
            (
                "ML Engineer, DataWorks (2020–2022)",
                [
                    "Trained classification models on imbalanced data with ROC-AUC 0.92.",
                    "Created data quality checks and monitoring dashboards.",
                ],
            ),
        ],
        "education": [
            "M.S. in Data Science, Tech Institute (2018–2020)",
        ],
    },
]

TXT_SAMPLE = """Name: Jordan Lee\nEmail: jordan.lee@example.com\nPhone: +1-555-3344\nLocation: Austin, TX\nLinkedIn: linkedin.com/in/jordanlee\n\nSummary\n-------\nFull-stack engineer focusing on React/Node and cloud-native deployments.\n\nSkills\n------\nJavaScript, TypeScript, React, Node.js, Express, PostgreSQL, Docker, Kubernetes\n\nExperience\n----------\nSenior Engineer, CloudNova (2022–Present)\n- Built React apps with hooks and Zustand; performance improved 25%.\n- Designed Node.js APIs and optimized SQL queries.\n\nEngineer, WebLift (2019–2022)\n- Migrated legacy jQuery app to React; reduced defects by 40%.\n\nEducation\n---------\nB.E. in Computer Engineering, City University\n"""


def create_docx(sample: dict, out_path: str) -> None:
    doc = Document()
    doc.add_heading(sample["full_name"], level=0)
    doc.add_paragraph(f"Email: {sample['email']}")
    doc.add_paragraph(f"Phone: {sample['phone']}")
    doc.add_paragraph(f"Location: {sample['location']}")

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(sample["summary"])

    doc.add_heading("Skills", level=1)
    doc.add_paragraph(", ".join(sample["skills"]))

    doc.add_heading("Experience", level=1)
    for role, bullets in sample["experience"]:
        doc.add_paragraph(role)
        for b in bullets:
            doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("Education", level=1)
    for edu in sample["education"]:
        doc.add_paragraph(edu)

    doc.save(out_path)


def main():
    here = os.path.dirname(os.path.abspath(__file__))
    for s in SAMPLES:
        out = os.path.join(here, s["name"])
        create_docx(s, out)
        print(f"Created {out}")

    txt_path = os.path.join(here, "sample_resume2.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(TXT_SAMPLE)
    print(f"Created {txt_path}")


if __name__ == "__main__":
    main()
